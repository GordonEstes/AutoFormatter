# -*- coding: utf-8 -*-

################
# BUGS:
    # initStyles does not properly initialize the font ("Arial" stays as "Calibri")
    # Other stuff
    # Other stuff
################

import sys
# sys.path.insert(0,"/Users/Gordon/Gordon's Files/AutoFormatter/lib")
import filelib, listlib, regexlib, doclib, numlib
import os, string
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import unicodedata
from docx.shared import Inches
from docx.shared import Pt
from story import Story
import Tkinter as tk

# The Formatter is an object that is able to take in a Story object and return a properly
# formatted version.
class Formatter(object):

    # Initializes the values used by the Formatter.
    def __init__(self):
        self.legalChars = "1234567890qwertyuiopasdfghjklzxcvbnmQWERTYUIOPASDFGHJKLZXCVBNM!#$%&*()_+=;:'\"/?.,—<>“”’ "
        self.story = None
        self.document1 = None
        self.document2 = None
        self.chapterNames = True
        self.progress = "Waiting for input."
        self.stage = "Incomplete"
        self.step = 0.0
        self.steps = 0.0
        self.gettingSize = False
        self.paragraphDict = {}
        self.numRuns = 0.0
        
    def initStyles(self):
        self.progress = "Initializing styles..."
        styles = self.document2.styles

        styles["Heading 1"].font.name = "Arial"
        styles["Heading 1"].font.size = Pt(16)
        styles["Heading 1"].font.color.rgb = None 
        styles["Heading 1"].font.bold = True
        styles["Heading 1"].paragraph_format.first_line_indent = Inches(0.5)

        styles["Heading 2"].font.name = "Arial"
        styles["Heading 2"].font.size = Pt(14)
        styles["Heading 2"].font.color.rgb = None 
        styles["Heading 2"].font.bold = True
        styles["Heading 2"].font.italic = True
        styles["Heading 2"].paragraph_format.first_line_indent = Inches(0.5)

        styles["Heading 3"].font.name = "Arial"
        styles["Heading 3"].font.size = Pt(13)
        styles["Heading 3"].font.color.rgb = None 
        styles["Heading 3"].font.bold = True
        styles["Heading 3"].paragraph_format.first_line_indent = Inches(0.5)

        styles["Normal"].font.name = "Times New Roman"
        styles["Normal"].font.size = Pt(14)
        styles["Normal"].font.color.rgb = None 
        styles["Normal"].paragraph_format.first_line_indent = Inches(0.5)

    # Takes in an object representing the story to be formatted and sets the local
    # story value to equal it.
    def take(self,story):
        self.progress = "Taking an unformatted story..."
        self.story = story
        self.document1 = Document(story.path)
        self.document2 = Document()
        self.initStyles()

    # Saves the final document to a docx with the '_formatted' suffix in the same directory as the
    # original.
    def save(self):
        self.progress = "Saving formatted document..."
        path = self.story.path
        document = self.document2
        path = regexlib.clipRight(path,".")
        path += "_formatted.docx"
        document.save(path)

    # Creates the document1 (original) and document2 (formatted) Document objects
    # and populates the latter using the plaintext version of the former. Also sets
    # the base formatting and style of document2.
    def build(self):
        self.progress = "Building document..."
        document1 = self.document1
        document2 = self.document2
        paragraph_format = document2.styles['Normal'].paragraph_format
        paragraph_format.space_before = 0 #Set paragraph spacing to 0 pica.
        paragraph_format.space_after = 0 
        for p1 in document1.paragraphs: #Copy each paragraph from document1 to document 2.
            p2 = document2.add_paragraph("")
            for r1 in p1.runs:
                r2 = p2.add_run(r1.text)
                r2.font.italic = r1.font.italic
                if "Italic" in r1.style.name: 
                    r2.font.italic = True
                self.numRuns += 1
        self.compressRuns()

    # Removes the following symbols from document2: "»","|","«","•","    "
    # Also changes em-dashes (—) into triple en-dashes (---) to avoid formatting
    # bugs later on.
    def removeSymbols(self):
        if not self.gettingSize: self.progress = "Removing symbols..."
        document2 = self.document2
        self.step = 0.0
        for paragraph in document2.paragraphs:
            for run in paragraph.runs:
                text = run.text
                text = text.replace(u"—","---")
                text = regexlib.removeAll(text,"»")
                text = regexlib.removeAll(text,"|")
                text = regexlib.removeAll(text,"«")
                text = regexlib.removeAll(text,"•")
                text = regexlib.removeAll(text,"    ")
                self.step += 1
                run.text = text

    # Inserts a properly formatted (Heading 1) title at the top of the document.
    def insertTitle(self):
        if not self.gettingSize: self.progress = self.progress = "Adding title..."
        title = self.story.title
        style = 'Heading 1'
        self.document2.paragraphs[0].insert_paragraph_before(title,style)

    # Inserts a properly formatted (Heading 2) author attribution beneath the document
    # title.
    def insertAuthor(self):
        if not self.gettingSize: self.progress = self.progress = "Adding author..."
        author = "by " + self.story.author
        style = 'Heading 2'
        self.document2.paragraphs[1].insert_paragraph_before(author,style)

    # Inserts the copyright date beneath the author, if applicable.
    def insertCopyright(self):
        if not self.gettingSize: self.progress = self.progress = "Adding copyright..."
        copyright = self.story.copyright
        if copyright == "": return
        paragraph = self.document2.paragraphs[2]
        copyrightParagraph = paragraph.insert_paragraph_before("")
        copyrightParagraph.add_run("(copyright %s)" % copyright).italic = True

    # Inserts the publisher, if applicable.
    def insertPublisher(self):
        if not self.gettingSize: self.progress = self.progress = "Adding publisher..."
        publisher = self.story.publisher
        if publisher == "": return
        paragraph = self.document2.paragraphs[2]
        publisherParagraph = paragraph.insert_paragraph_before("")
        publisherParagraph.add_run("(publisher %s)" % publisher).italic = True

    # Finds each chapter header and formats it appropriately (i.e., "Chapter __" becomes
    # Heading 2, and the chapter names, if present, become Heading 3)
    def formatChapters(self):
        if not self.gettingSize: self.progress = "Formatting chapters..."
        chapter = False
        for paragraph in self.document2.paragraphs:
            text = paragraph.text
            if "CHAPTER" in text or "Chapter" in text or "BOOK" in text or "Book" in text:
                if len(text) <= len("Chapter XXXXIIII"):
                   chapter = True
                   paragraph.style = 'Heading 2'
            elif chapter and self.chapterNames:
                paragraph.style = 'Heading 3'
                chapter = False
            else:
                chapter = False
            self.step += 1

    # Removes the scanner error of mashing two separate lines of dialogue together into one
    # paragraph.
    def fixDoubleQuotes(self):
        if not self.gettingSize: self.progress = "Fixing double quotes..."
        self.psteps = 0.0
        for paragraph in self.document2.paragraphs:
            text = paragraph.text
            i = regexlib.match(text,u'” “')  
            if i == -1: i = regexlib.match(text,'" "')      
            if i != -1:
                text1 = text[:i+1]
                text2 = text[i+2:]
                paragraph.text = text2
                paragraph.insert_paragraph_before(text1)
            self.step += 1

    # Translates unicode punctuation to ASCII punctuation.
    def convertPunctuation(self):
        if not self.gettingSize: self.progress = "Converting punctuation..."
        for paragraph in self.document2.paragraphs:
            for run in paragraph.runs:
                text = run.text
                punctuation = { 0x2018:0x27, 0x2019:0x27, 0x201C:0x22, 0x201D:0x22 }
                text = u'%s' % text
                text = text.translate(punctuation).encode('ascii', 'ignore')
                run.text = text
                self.step += 1

    # Replaces '---' with '—' and fixes em-dash spacing errors. 
    def fixEmDash(self):
        if not self.gettingSize: self.progress = "Fixing em-dashes..."
        document2 = self.document2
        for paragraph in document2.paragraphs:
            for run in paragraph.runs:
                text = run.text
                text = text.replace('---"',u'---”')
                text = text.replace('"---',u'“---')
                text = text.replace("---",u"—")
                text = regexlib.replaceSub(text," — ","—")
                text = regexlib.replaceSub(text," —","—")
                text = regexlib.replaceSub(text,"— ","—")
                run.text = text
                self.step += 1

    # Fixes some common quotation mark errors and changes ASCII quotation marks
    # to the appropriate Unicode quotation marks.
    def fixQuotations(self):
        if not self.gettingSize: self.progress = "Fixing quotations..."
        for paragraph in self.document2.paragraphs:
            for run in paragraph.runs:
                text = run.text
                text = regexlib.replaceSub(text,' "',u' “')
                text = regexlib.replaceSub(text,'" ',u'” ')
                text = regexlib.replaceSub(text,'."',u'.”')
                text = regexlib.replaceSub(text,',"',u',”')
                text = regexlib.replaceSub(text,'!"',u'!”')
                text = regexlib.replaceSub(text,'?"',u'?”')
                text = regexlib.replaceSub(text,'..."',u'...”')
                text = regexlib.replaceSub(text,'#—"#',u'#—“#')
                text = regexlib.replaceSub(text,'#"—#',u'#”—#')
                text = regexlib.replaceSub(text,'#—"',u'#—”')
                text = regexlib.replaceSub(text,'"—#',u'“—#')
                text = regexlib.replaceSub(text,'’"',u'’”')
                text = regexlib.replaceSub(text,'\'"',u'’”')
                text = regexlib.replaceSub(text,'";',u'”;')
                text = regexlib.replaceSub(text,'":',u'”:')
                text = regexlib.replaceSub(text,"' \"",u'’”')
                text = regexlib.replaceSub(text,"’ \"",u'’”')
                text = regexlib.replaceSub(text,'"',u'“')
                run.text = text
                self.step += 1

    # Removes the all-caps word(s) beginning each chapter, which are common to 
    # older stories
    def fixCaps(self):
        if not self.gettingSize: 
            self.progress = "Fixing chapter capitalization..."
        for paragraph in self.document2.paragraphs:
            if len(paragraph.text) < 3: 
                continue
            if paragraph.style.name != "Normal":
                continue
            if (paragraph.text[0] in string.ascii_uppercase
            and paragraph.text[1] in string.ascii_uppercase
            and paragraph.text[2] in string.ascii_uppercase):
                s = paragraph.text
                for i in xrange(1,len(s)):
                    if s[i] in string.ascii_lowercase: break
                    if s[i] in string.ascii_uppercase:
                        s = regexlib.replaceIndex(s,i,string.lower(s[i]))
                paragraph.text = s
            self.step += 1

    def mergeParagraphs(self,p1,p2):
        for r2 in p2.runs:
            r1 = p1.add_run(r2.text)
            r1.font.italic = r2.font.italic
            r2.clear()
        self.deleteParagraph(p2)

    def deleteParagraph(self,paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Finds paragraphs beginning with a lowercase character and re-attaches them to the
    # end of the preceding paragraph.
    def fixCarriageReturn(self):
        if not self.gettingSize: self.progress = "Fixing carriage returns..."
        last = None
        for paragraph in self.document2.paragraphs:
            if paragraph == self.document2.paragraphs[0]: continue
            text = paragraph.text
            if last != None and len(text) >= 1 and text[0] in " qwertyuiopasdfghjklzxcvbnm":
                self.mergeParagraphs(last,paragraph)
                # doclib.mergeParagraphs(last,paragraph)
            else:
                last = paragraph
            self.step += 1

    # Fixes some common apostrophe errors and changes ASCII apostrophes to the
    # appropriate Unicode apostrophe.
    def fixApostrophes(self):
        if not self.gettingSize: self.progress = "Fixing apostrophes..."
        for paragraph in self.document2.paragraphs:
            for run in paragraph.runs:
                text = run.text
                text = regexlib.replaceSub(text,"''","\"")
                text = regexlib.replaceSub(text," '",u' ‘')
                text = regexlib.replaceSub(text,"' ",u'’ ')
                text = regexlib.replaceSub(text,".'",u'.’')
                text = regexlib.replaceSub(text,",'",u',’')
                text = regexlib.replaceSub(text,"!'",u'!’')
                text = regexlib.replaceSub(text,"?'",u'?’')
                text = regexlib.replaceSub(text,"...'",u'...’')
                text = regexlib.replaceSub(text,"'.",u'’.')
                text = regexlib.replaceSub(text,"#—'#",u'#—‘#')
                text = regexlib.replaceSub(text,"#'—#",u'#’—#')
                text = regexlib.replaceSub(text,"#—'",u'#—’')
                text = regexlib.replaceSub(text,"'—#",u'‘—#')
                text = regexlib.replaceSub(text,"',",u'’,')
                text = regexlib.replaceSub(text,"'!",u'’!')
                text = regexlib.replaceSub(text,"'?",u'’?')
                text = regexlib.replaceSub(text,"'...",u'’...')
                text = regexlib.replaceSub(text,"#'#",u'#’#')
                text = regexlib.replaceSub(text,"'",u'‘')
                run.text = text
                self.step += 1

    # Removes double spaces and replaces them with single spaces.
    def fixDoubleSpace(self):
        if not self.gettingSize: self.progress = "Fixing double spaces..."
        for paragraph in self.document2.paragraphs:
            for run in paragraph.runs:
                while regexlib.match(run.text, "  ") != -1:
                    run.text = regexlib.replaceSub(run.text, "  ", " ")
                self.step += 1

    # Changes common ellipse misprints to proper formatting.
    def fixEllipses(self):
        if not self.gettingSize: self.progress = "Fixing ellipses..."
        for paragraph in self.document2.paragraphs:
            for run in paragraph.runs:
                text = run.text
                text = regexlib.replaceSub(text," . . . . .", "...")
                text = regexlib.replaceSub(text," . . . .","...")
                text = regexlib.replaceSub(text," . . .","...")
                text = regexlib.replaceSub(text,". . . . .", "...")
                text = regexlib.replaceSub(text,". . . .","...")
                text = regexlib.replaceSub(text,". . .","...")
                text = regexlib.replaceSub(text,".....","...")
                text = regexlib.replaceSub(text,"....","...")
                text = regexlib.replaceSub(text," . .",u'...”')
                text = regexlib.replaceSub(text,". . ",u"“...")
                run.text = text   
                self.step += 1

    # Fixes hyphen spacing, removes nonnecessary asterisks, and removes
    # additional symbols.      
    def fixPunctuation(self):
        if not self.gettingSize: self.progress = "Fixing punctuation..."
        for paragraph in self.document2.paragraphs:
            for run in paragraph.runs:
                text = run.text
                #Hyphens
                text = regexlib.removeSub(text," -")
                text = regexlib.replaceSub(text,"- ","-")
                #Misc
                text = regexlib.replaceSub(text,"/","I")
                text = regexlib.removeSub(text,"\\")
                text = regexlib.removeSub(text,"^")
                #Asterisks
                text = regexlib.replaceSub(text,"* * * *", "& & & &")
                text = regexlib.removeSub(text,"*")
                text = regexlib.replaceSub(text,"& & & &", "* * * *")
                # Periods and Spacing
                text = regexlib.replaceSub(text,"# .#","#  #")
                text = regexlib.replaceSub(text,"  "," ")
                text = regexlib.replaceSub(text,"“ ‘","“‘")
                run.text = text
                self.step += 1

    # Replaces/fixes words commonly mistaken by the scanner for one another.
    def fixWords(self):
        if not self.gettingSize: self.progress = "Fixing words..."
        for paragraph in self.document2.paragraphs:
            for run in paragraph.runs:
                text = run.text
                text = regexlib.replaceWord(text,"comer","corner")
                text = regexlib.replaceWord(text,"bom","born")
                text = regexlib.replaceWord(text,"modem","modern")
                text = regexlib.replaceWord(text,"tiling","thing")
                text = regexlib.replaceWord(text,"diat","that")
                text = regexlib.replaceWord(text,"sec","see")
                text = regexlib.replaceWord(text,"secs","sees")
                text = regexlib.replaceWord(text,"Fd","I'd")
                text = regexlib.replaceWord(text,"diem","them")
                text = regexlib.replaceWord(text,"Modem","Modern")
                text = regexlib.replaceSub(text,"‘Tm","“I'm")
                text = regexlib.replaceWord(text,"tire","the")
                text = regexlib.replaceSub(text,"boy friend","boyfriend")
                text = regexlib.replaceSub(text,"girl friend","girlfriend")
                text = regexlib.replaceSub(text,"Pie ", "He")
                text = regexlib.replaceSub(text,"Fie,", "He")
                run.text = text
                self.step += 1

    # Retrieves the total number of steps to be taken within the document in order
    # to properly set the loading bar.
    def getSize(self):
        self.progress = "Getting document size..."
        document2 = self.document2
        self.document2 = Document()
        self.steps = 19*self.numRuns
        for paragraph in self.document1.paragraphs: 
            p = self.document2.add_paragraph("")
            for run in paragraph.runs:
                r = p.add_run("")
        self.gettingSize = True
        self.format()
        self.fix()
        self.gettingSize = False
        self.steps = self.step
        self.step = 0.0
        self.document2 = document2

    # Spaces and justifies the "* * * *" scene break common to stories.
    def formatSceneBreaks(self):
        if not self.gettingSize: self.progress = "Separating scenes..."
        for paragraph in self.document2.paragraphs:
            if paragraph.text == "* * * *":
                text = ""
                style = "Normal"
                paragraph_1 = paragraph.insert_paragraph_before(text,style)
                paragraph_2 = paragraph.insert_paragraph_before(paragraph.text,style)
                paragraph_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.text = ""
            self.step += 1

    def fixPeriodSpacing(self):
        if not self.gettingSize: self.progress = "Fixing period spacing..."
        for paragraph in self.document2.paragraphs:
            for run in paragraph.runs:
                text = run.text
                while (True):
                    k = regexlib.match(text,"#. <")
                    if k == -1: break
                    new = text[:k+1]
                    new += text[k+2:]
                    text = new
                run.text = text
                self.step += 1
                while (True):
                    k = regexlib.match(text,"# .<")
                    if k == -1: break
                    new = text[:k+1]
                    new += text[k+2:]
                    text = new
                run.text = text
                self.step += 1
                while (True):
                    k = regexlib.match(text,"# . >")
                    if k == -1: break
                    new = text[:k+1]
                    new += text[k+3:]
                    text = new
                run.text = text 
                self.step += 1 

    def compressRuns(self):
        self.progress = "Compressing runs..."
        for paragraph in self.document2.paragraphs:
            i = 0
            while i < len(paragraph.runs)-1:
                r1 = paragraph.runs[i]
                r2 = paragraph.runs[i+1]
                if r1.font.italic == r2.font.italic:
                    r1.text = r1.text + r2.text
                    r2.clear()
                i += 1


    # Calls the other methods in their proper order.
    def fix(self):
        n = 0
        if not self.gettingSize: self.progress = "Fixing mistakes..."
        self.fixEmDash()
        self.fixDoubleQuotes()
        self.fixApostrophes()
        self.fixQuotations()
        self.fixDoubleSpace()
        self.fixEllipses()
        self.fixPunctuation()
        self.fixWords()
        self.fixCaps()
        self.fixDoubleQuotes()
        self.fixCarriageReturn()
        self.formatSceneBreaks()
        self.insertTitle()
        self.insertAuthor()
        self.insertCopyright()
        self.insertPublisher()
        self.fixPeriodSpacing()

    # Calls the other methods in their proper order.
    def format(self):
        if not self.gettingSize: self.progress = "Formatting story..."
        self.removeSymbols()
        self.formatChapters()

    # Opens the formatted docx after it's been saved.
    def open(self):
        self.progress = "Opening file..."
        path = self.story.path
        path = regexlib.clipRight(path,".")
        path += "_formatted.docx"
        os.system("start "+path)

    # Sets the local result variable by processing the local story variable.
    def run(self):
        self.progress = "Running formatter..."
        sys.stdout.flush()
        self.build()
        self.getSize()
        self.format()
        self.fix()
        self.save()
        self.stage = "Complete"
        self.open()
        self.step = 0.0
        self.steps = 0.0
